from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from models import Assignment, Problem, ProblemImage

OPTION_LINE_RE = re.compile(r"^[A-DFTＡ-ＤＦＴ][.．、]\s+")
TRUE_FALSE_LINE_RE = re.compile(r"^[TFＴＦ]$", re.IGNORECASE)


class DocxWriter:
    def write_document(
        self,
        course_name: str,
        assignments: list[Assignment],
        output_dir: Path,
        *,
        filename_stem: str | None = None,
    ) -> Path:
        output_dir.mkdir(parents=True, exist_ok=True)
        document = Document()
        self._configure_page(document)
        self._configure_styles(document)

        for index, assignment in enumerate(assignments):
            if index > 0:
                document.add_section(WD_SECTION.NEW_PAGE)
            self._add_assignment(document, assignment)

        if filename_stem:
            filename = f"{self._safe_name(filename_stem)}.docx"
        else:
            filename = f"{self._safe_name(course_name)}.docx"
        target = output_dir / filename
        document.save(target)
        return target

    def _configure_page(self, document: Document) -> None:
        section = document.sections[0]
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    def _configure_styles(self, document: Document) -> None:
        normal = document.styles["Normal"]
        normal.font.name = "宋体"
        self._set_rfonts(normal._element.rPr.rFonts, "宋体")
        normal.font.size = Pt(11)
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        normal.paragraph_format.space_after = Pt(6)

        for style_name, font_size, color in (
            ("Heading 1", 16, RGBColor(31, 60, 115)),
            ("Heading 2", 13, RGBColor(48, 48, 48)),
            ("Heading 3", 11.5, RGBColor(85, 85, 85)),
        ):
            style = document.styles[style_name]
            self._set_heading_style_fonts(style)
            style.font.size = Pt(font_size)
            style.font.bold = True
            style.font.color.rgb = color
            style.paragraph_format.space_before = Pt(10 if style_name != "Heading 3" else 8)
            style.paragraph_format.space_after = Pt(4)

        if "CodeBlock" not in document.styles:
            style = document.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = "Consolas"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
            style.font.size = Pt(10)
            style.paragraph_format.space_after = Pt(0)
            style.paragraph_format.left_indent = Pt(2)

    def _add_assignment(self, document: Document, assignment: Assignment) -> None:
        heading = document.add_heading(assignment.title, level=1)
        heading.paragraph_format.keep_with_next = True
        self._apply_heading_run_fonts(heading)

        if assignment.warnings:
            warning = document.add_paragraph()
            warning.paragraph_format.space_after = Pt(8)
            run = warning.add_run("抓取提醒：")
            run.bold = True
            run.font.color.rgb = RGBColor(163, 93, 15)
            warning.add_run("；".join(assignment.warnings))

        for problem in assignment.problems:
            self._add_problem(document, problem)

    def _add_problem(self, document: Document, problem: Problem) -> None:
        heading = document.add_heading(problem.title, level=2)
        heading.paragraph_format.keep_with_next = True
        self._apply_heading_run_fonts(heading)

        for section in problem.sections:
            if problem.samples and section.kind.startswith("sample"):
                continue
            show_heading = not (section.kind == "description" and section.title == "题目描述")
            if show_heading:
                section_heading = document.add_paragraph(section.title, style="Heading 3")
                section_heading.paragraph_format.keep_with_next = True
                self._apply_heading_run_fonts(section_heading)
            for block in [piece.strip() for piece in section.content.split("\n\n") if piece.strip()]:
                if section.kind.startswith("sample"):
                    self._add_code_table(document, block)
                else:
                    self._add_body_content(document, block)

        if problem.samples:
            sample_heading = document.add_paragraph("样例", style="Heading 3")
            sample_heading.paragraph_format.keep_with_next = True
            self._apply_heading_run_fonts(sample_heading)
            for index, sample in enumerate(problem.samples, start=1):
                if sample.input_text:
                    document.add_paragraph(f"样例输入 {index}")
                    self._add_code_table(document, sample.input_text)
                if sample.output_text:
                    document.add_paragraph(f"样例输出 {index}")
                    self._add_code_table(document, sample.output_text)
                if sample.note:
                    self._add_body_paragraph(document, sample.note)

        self._add_images(document, problem.images)
        document.add_paragraph()

    def _add_body_paragraph(self, document: Document, content: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Pt(18)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.add_run(content)

    def _add_body_content(self, document: Document, content: str) -> None:
        option_lines = [line.strip() for line in content.splitlines() if line.strip()]
        if option_lines and all(TRUE_FALSE_LINE_RE.fullmatch(line) for line in option_lines):
            self._add_true_false_lines(document, option_lines)
            return
        if option_lines and all(OPTION_LINE_RE.match(line) for line in option_lines):
            for option_line in option_lines:
                self._add_option_paragraph(document, option_line)
            return
        self._add_body_paragraph(document, content)

    def _add_option_paragraph(self, document: Document, content: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Pt(18)
        paragraph.paragraph_format.first_line_indent = Pt(-12)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.add_run(content)

    def _add_true_false_lines(self, document: Document, lines: list[str]) -> None:
        for index, line in enumerate(lines):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Pt(18)
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            paragraph.paragraph_format.space_after = Pt(0 if index < len(lines) - 1 else 6)
            paragraph.add_run(line)

    def _add_code_table(self, document: Document, content: str) -> None:
        table = document.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        table.autofit = True
        cell = table.rows[0].cells[0]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        self._shade_cell(cell, "F6F8FB")
        paragraph = cell.paragraphs[0]
        paragraph.style = "CodeBlock"
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.right_indent = Pt(0)
        run = paragraph.add_run(content)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
        run.font.size = Pt(10)

    def _add_images(self, document: Document, images: list[ProblemImage]) -> None:
        valid_images = [image for image in images if image.local_path and Path(image.local_path).exists()]
        if not valid_images:
            return

        image_heading = document.add_paragraph("题面图片", style="Heading 3")
        self._apply_heading_run_fonts(image_heading)
        for image in valid_images:
            path = Path(image.local_path)
            picture_paragraph = document.add_paragraph()
            picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            picture_run = picture_paragraph.add_run()
            picture_run.add_picture(str(path), width=Inches(5.9))
            if image.alt and not re.fullmatch(r"[A-Za-z0-9_.-]+\.(png|jpe?g|gif|webp|svg)", image.alt, re.IGNORECASE):
                caption = document.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.paragraph_format.space_after = Pt(8)
                caption_run = caption.add_run(image.alt)
                caption_run.font.color.rgb = RGBColor(110, 110, 110)

    def _shade_cell(self, cell, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)

    def _set_heading_style_fonts(self, style) -> None:
        style.font.name = "宋体"
        r_fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
        self._set_rfonts(r_fonts, "宋体")

    def _set_heading_run_fonts(self, run) -> None:
        run.bold = True
        run.font.name = "宋体"
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.get_or_add_rFonts()
        self._set_rfonts(r_fonts, "宋体")

    def _apply_heading_run_fonts(self, paragraph) -> None:
        for run in paragraph.runs:
            self._set_heading_run_fonts(run)

    def _set_rfonts(self, r_fonts, font_name: str) -> None:
        for theme_attr in ("w:asciiTheme", "w:hAnsiTheme", "w:cstheme", "w:eastAsiaTheme"):
            key = qn(theme_attr)
            if key in r_fonts.attrib:
                del r_fonts.attrib[key]
        r_fonts.set(qn("w:ascii"), font_name)
        r_fonts.set(qn("w:hAnsi"), font_name)
        r_fonts.set(qn("w:cs"), font_name)
        r_fonts.set(qn("w:eastAsia"), font_name)

    def _safe_name(self, value: str) -> str:
        cleaned = re.sub(r'[\\/:*?"<>|]+', "_", value).strip()
        return cleaned or "PTA"
