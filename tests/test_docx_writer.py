from __future__ import annotations

import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document
from PIL import Image

from export.docx_writer import DocxWriter
from models import Assignment, Problem, ProblemImage, ProblemSample, ProblemSection


class DocxWriterTests(unittest.TestCase):
    def test_generates_openable_docx_with_expected_sections(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            image_path = temp_path / "sample.png"
            Image.new("RGB", (40, 40), color=(255, 80, 80)).save(image_path)

            assignment = Assignment(
                id="a1",
                title="第一次作业",
                url="https://pintia.cn/problem-sets/homework-1",
                course_name="PTA题目集",
                problems=[
                    Problem(
                        id="p1",
                        title="L1-001 Hello PTA",
                        url="https://pintia.cn/problem-sets/homework-1/problems/l1-001",
                        score="25",
                        sequence_label="L1-001",
                        title_source="list-link",
                        sections=[
                            ProblemSection(kind="description", title="题目描述", content="请输出 Hello PTA。"),
                            ProblemSection(kind="input", title="输入格式", content="本题没有输入。"),
                            ProblemSection(kind="output", title="输出格式", content="输出一行 Hello PTA。"),
                        ],
                        samples=[ProblemSample(input_text="", output_text="Hello PTA")],
                        images=[ProblemImage(url="https://example.com/image.png", alt="示意图", local_path=str(image_path))],
                    )
                ],
            )

            output = DocxWriter().write_document("PTA题目集", [assignment], temp_path)

            self.assertTrue(output.exists())
            opened = Document(str(output))
            text = "\n".join(paragraph.text for paragraph in opened.paragraphs)
            self.assertIn("第一次作业", text)
            self.assertIn("L1-001 Hello PTA", text)
            self.assertNotIn("25 分", text)
            self.assertNotIn("题目链接", text)
            non_empty_paragraphs = [paragraph.text for paragraph in opened.paragraphs if paragraph.text.strip()]
            self.assertEqual("第一次作业", non_empty_paragraphs[0])
            self.assertGreaterEqual(len(opened.inline_shapes), 1)

    def test_ignores_filename_like_image_caption(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            image_path = temp_path / "sample.png"
            Image.new("RGB", (40, 40), color=(255, 80, 80)).save(image_path)

            assignment = Assignment(
                id="a1",
                title="第一次作业",
                url="https://pintia.cn/problem-sets/homework-1",
                problems=[
                    Problem(
                        id="p1",
                        title="题目一",
                        url="https://pintia.cn/problem-sets/homework-1/problems/l1-001",
                        images=[ProblemImage(url="https://example.com/image.png", alt="image.png", local_path=str(image_path))],
                    )
                ],
            )

            output = DocxWriter().write_document("PTA题目集", [assignment], temp_path)
            opened = Document(str(output))
            text = "\n".join(paragraph.text for paragraph in opened.paragraphs)
            self.assertNotIn("image.png", text)

    def test_supports_custom_filename_stem(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            assignment = Assignment(
                id="a1",
                title="Homework 1",
                url="https://pintia.cn/problem-sets/homework-1",
                problems=[Problem(id="p1", title="Problem 1", url="https://pintia.cn/problems/p1")],
            )

            output = DocxWriter().write_document(
                "PTA题目集",
                [assignment],
                temp_path,
                filename_stem="01_homework_1",
            )

            self.assertTrue(output.exists())
            self.assertEqual("01_homework_1.docx", output.name)

    def test_hides_description_heading_and_keeps_options_inline(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            assignment = Assignment(
                id="a1",
                title="第一次作业",
                url="https://pintia.cn/problem-sets/homework-1",
                problems=[
                    Problem(
                        id="p1",
                        title="2-1",
                        url="https://pintia.cn/problem-sets/homework-1/problems/2-1",
                        sections=[
                            ProblemSection(
                                kind="description",
                                title="题目描述",
                                content="1-1 因特网的前身是 1969 年创建的第一个分组交换网（）。\nA. internet\nB. Internet\nC. NSFNET\nD. ARPANET",
                            )
                        ],
                    )
                ],
            )

            output = DocxWriter().write_document("PTA题目集", [assignment], temp_path)
            opened = Document(str(output))
            paragraphs = [paragraph.text for paragraph in opened.paragraphs if paragraph.text.strip()]
            self.assertNotIn("题目描述", paragraphs)
            merged_body = next(paragraph for paragraph in paragraphs if "A. internet" in paragraph)
            self.assertIn("A. internet", merged_body)
            self.assertIn("B. Internet", merged_body)

    def test_all_option_lines_use_same_option_indent(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            assignment = Assignment(
                id="a1",
                title="第一次作业",
                url="https://pintia.cn/problem-sets/homework-1",
                problems=[
                    Problem(
                        id="p1",
                        title="2-77 6-1 在客户/服务器模型中，客户指的是（）。",
                        url="https://pintia.cn/problem-sets/homework-1/problems/2-77",
                        sections=[
                            ProblemSection(
                                kind="description",
                                title="题目描述",
                                content="A. 请求方\nB. 响应方\nC. 硬件\nD. 软件",
                            )
                        ],
                    )
                ],
            )

            output = DocxWriter().write_document("PTA题目集", [assignment], temp_path)
            opened = Document(str(output))
            option_paragraphs = [paragraph for paragraph in opened.paragraphs if paragraph.text in {"A. 请求方", "B. 响应方", "C. 硬件", "D. 软件"}]
            self.assertEqual(4, len(option_paragraphs))
            first = option_paragraphs[0].paragraph_format
            for paragraph in option_paragraphs[1:]:
                current = paragraph.paragraph_format
                self.assertEqual(first.left_indent, current.left_indent)
                self.assertEqual(first.first_line_indent, current.first_line_indent)

    def test_formats_t_f_options_and_preserves_fill_blank_text(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            assignment = Assignment(
                id="a1",
                title="题型回归",
                url="https://pintia.cn/problem-sets/homework-1",
                problems=[
                    Problem(
                        id="p1",
                        title="判断题",
                        url="https://pintia.cn/problems/judge",
                        sections=[
                            ProblemSection(
                                kind="description",
                                title="题目描述",
                                content="TCP 是面向连接的传输层协议。\nT. 正确\nF. 错误",
                            )
                        ],
                    ),
                    Problem(
                        id="p2",
                        title="填空题",
                        url="https://pintia.cn/problems/blank",
                        sections=[
                            ProblemSection(
                                kind="description",
                                title="题目描述",
                                content="请补全 C 语言中的标准输出语句：printf(    );\n函数名是 ________。",
                            )
                        ],
                    ),
                ],
            )

            output = DocxWriter().write_document("PTA题目集", [assignment], temp_path)
            opened = Document(str(output))
            paragraphs = [paragraph.text for paragraph in opened.paragraphs if paragraph.text.strip()]
            merged_text = "\n".join(paragraphs)

            self.assertIn("T. 正确", merged_text)
            self.assertIn("F. 错误", merged_text)
            self.assertIn("请补全 C 语言中的标准输出语句：printf(    );", merged_text)


    def test_heading_style_uses_bold_songti_for_mixed_title(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            assignment = Assignment(
                id="a1",
                title="作业标题",
                url="https://pintia.cn/problem-sets/homework-1",
                problems=[
                    Problem(
                        id="p1",
                        title="2-1 6-1 在 DNS 的递归查询中，给客户端返回地址的是（）。",
                        url="https://pintia.cn/problem-sets/homework-1/problems/2-1",
                    )
                ],
            )
            output = DocxWriter().write_document("PTA题目集", [assignment], temp_path)
            with zipfile.ZipFile(output) as archive:
                document_xml = archive.read("word/document.xml").decode("utf-8")
                styles_xml = archive.read("word/styles.xml").decode("utf-8")

            self.assertIn('w:styleId="Heading2"', styles_xml)
            self.assertIn('w:ascii="宋体"', styles_xml)
            self.assertIn('w:hAnsi="宋体"', styles_xml)
            self.assertIn('w:eastAsia="宋体"', styles_xml)
            self.assertIn("<w:b/>", styles_xml)

            title_index = document_xml.index(assignment.problems[0].title)
            title_slice = document_xml[max(0, title_index - 300) : title_index + 300]
            self.assertIn('w:rFonts w:ascii="宋体"', title_slice)
            self.assertIn('w:hAnsi="宋体"', title_slice)
            self.assertIn('w:eastAsia="宋体"', title_slice)
            self.assertIn("<w:b/>", title_slice)


if __name__ == "__main__":
    unittest.main()
